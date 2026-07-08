"""
文档处理服务

处理文档上传、文本提取、智能分段等核心逻辑
支持三层知识库导航：标题索引 → 段落 → 相邻段落
"""

import asyncio
import os
import re
import hashlib
from typing import List, Dict, Any

import markdown_it
from fastapi import UploadFile

from app.config.settings import settings


def _write_bytes(path, content: bytes) -> None:
    """同步写入字节到文件（供 asyncio.to_thread 调用）"""
    with open(path, "wb") as f:
        f.write(content)


class DocumentProcessor:
    """
    文档处理器

    支持:
    - txt 文本文件提取
    - docx Word 文档提取
    - 三层结构化分段（标题索引 + 细粒度段落 + 相邻关系）
    """

    KNOWLEDGE_UPLOAD_SUBDIR = "knowledge"
    MAX_SEGMENT_SIZE = 1000
    MIN_SEGMENT_SIZE = 50

    TITLE_PATTERNS = [
        (re.compile(r"^(#{1,6})\s+.+"), "md"),
        (re.compile(r"^([一二三四五六七八九十]+)[、.．].+"), "cn_num"),
        (re.compile(r"^(\d+)[、.．]\s*.+"), "num"),
        (re.compile(r"^\(([一二三四五六七八九十]+)\).+"), "cn_paren"),
        (re.compile(r"^\((\d+)\).+"), "num_paren"),
        (re.compile(r"^第([一二三四五六七八九十\d]+)[章节条].+"), "chapter"),
    ]

    def __init__(self):
        self._upload_dir = settings.get_absolute_path(
            f"{settings.upload_dir}/{self.KNOWLEDGE_UPLOAD_SUBDIR}"
        )
        self._upload_dir.mkdir(parents=True, exist_ok=True)

    async def extract_text(self, file: UploadFile) -> str:
        """
        从上传文件中提取文本内容

        Args:
            file: 上传的文件对象

        Returns:
            提取的文本内容

        Raises:
            ValueError: 不支持的文件类型
        """
        filename = file.filename or ""
        ext = self._get_file_extension(filename)

        content = await file.read()

        if ext == "txt":
            return self._extract_from_txt(content)
        elif ext == "docx":
            return self._extract_from_docx(content)
        elif ext == "xlsx":
            return self._extract_from_xlsx(content)
        elif ext == "md":
            return self._extract_from_txt(content)
        else:
            raise ValueError(f"不支持的文件类型: {ext}，仅支持 txt、md、docx、xlsx")

    def extract_text_from_path(self, file_path: str, file_type: str) -> str:
        """
        从磁盘文件路径提取文本内容（用于定时任务异步处理）

        Args:
            file_path: 文件绝对路径
            file_type: 文件类型（txt/md/docx/pdf）

        Returns:
            提取的文本内容

        Raises:
            ValueError: 不支持的文件类型或文件不存在
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        with open(file_path, "rb") as f:
            content = f.read()

        if file_type in ("txt", "md"):
            return self._extract_from_txt(content)
        elif file_type == "docx":
            return self._extract_from_docx(content)
        elif file_type == "xlsx":
            return self._extract_from_xlsx(content)
        elif file_type == "pdf":
            from app.utils.pdf_parser import pdf_parser

            return pdf_parser.parse_to_markdown(file_path)
        else:
            raise ValueError(
                f"不支持的文件类型: {file_type}，仅支持 txt、md、docx、xlsx、pdf"
            )

    def _get_file_extension(self, filename: str) -> str:
        """获取文件扩展名"""
        return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    def _extract_from_txt(self, content: bytes) -> str:
        """从 txt 文件提取文本"""
        encodings = ["utf-8", "gbk", "gb2312", "utf-16"]

        for encoding in encodings:
            try:
                return content.decode(encoding)
            except (UnicodeDecodeError, LookupError):
                continue

        raise ValueError("无法识别文件编码，请使用 UTF-8 编码")

    def _extract_from_docx(self, content: bytes) -> str:
        """从 docx 文件提取文本"""
        import io
        from docx import Document

        doc = Document(io.BytesIO(content))

        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style_name = para.style.name if para.style else ""
                if "Heading" in style_name or "Title" in style_name:
                    level = self._get_heading_level(style_name)
                    paragraphs.append(f"{'#' * level} {text}")
                else:
                    paragraphs.append(text)

        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip(" |"):
                    table_text.append(row_text)
            if table_text:
                paragraphs.append("\n".join(table_text))

        return "\n\n".join(paragraphs)

    def _get_heading_level(self, style_name: str) -> int:
        """从样式名称获取标题级别"""
        match = re.search(r"Heading\s*(\d+)", style_name)
        if match:
            return int(match.group(1))
        if "Title" in style_name:
            return 1
        return 2

    def _extract_from_xlsx(self, content: bytes) -> str:
        """从 xlsx 文件提取文本，每个 sheet 转为 markdown 表格，合并单元格值填充到所有子格"""
        import io
        from openpyxl import load_workbook

        with load_workbook(io.BytesIO(content), read_only=False, data_only=True) as wb:
            sheet_texts = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]

                if ws.max_row is None or ws.max_row < 1:
                    continue

                merged_values = self._resolve_merged_cells(ws)

                table_lines = []
                for row in ws.iter_rows(
                    min_row=1, max_row=ws.max_row, max_col=ws.max_column
                ):
                    cells = []
                    for cell in row:
                        val = merged_values.get((cell.row, cell.column), cell.value)
                        if val is None:
                            cells.append("")
                        else:
                            cells.append(str(val))
                    table_lines.append("| " + " | ".join(cells) + " |")

                if table_lines:
                    col_count = max(len(line.split("|")) - 2 for line in table_lines)
                    separator = "| " + " | ".join(["---"] * col_count) + " |"
                    md_table = f"## {sheet_name}\n\n{table_lines[0]}\n{separator}\n"
                    md_table += "\n".join(table_lines[1:])
                    sheet_texts.append(md_table)

        if not sheet_texts:
            return ""

        return "\n\n".join(sheet_texts)

    @staticmethod
    def _resolve_merged_cells(ws) -> dict:
        """
        解析工作表的合并单元格，将左上角值填充到合并范围内的所有单元格

        Args:
            ws: openpyxl 工作表

        Returns:
            {(row, col): value} 映射
        """
        merged_values = {}
        for merged_range in ws.merged_cells.ranges:
            top_left_val = ws.cell(merged_range.min_row, merged_range.min_col).value
            if top_left_val is None:
                continue
            for row in range(merged_range.min_row, merged_range.max_row + 1):
                for col in range(merged_range.min_col, merged_range.max_col + 1):
                    merged_values[(row, col)] = top_left_val
        return merged_values

    # ---- 三层结构化分段 ----

    def smart_segment(self, text: str) -> Dict[str, Any]:
        """
        三层结构化分段

        提取标题层级结构，在标题区域内进行细粒度段落切分，
        段落通过 segment_index 建立相邻关系，通过 title_id 关联标题

        Args:
            text: 原始文本

        Returns:
            {"titles": [...], "segments": [...]}
            titles: [{title_index, level, title, start_segment_index, end_segment_index}, ...]
            segments: [{segment_index, title, title_index, content, word_count}, ...]
        """
        title_entries = self._extract_title_structure(text)

        if not title_entries:
            segments = self._split_paragraphs_fine_grained(
                text.strip(), title_text="", title_index=-1
            )
            for i, seg in enumerate(segments):
                seg["segment_index"] = i
                seg["word_count"] = len(seg["content"])
            return {"titles": [], "segments": segments}

        titles = []
        all_segments = []
        segment_counter = 0

        for entry in title_entries:
            section_text = entry["content"]
            if not section_text.strip():
                continue

            title_text = entry["title"]
            title_index = entry["title_index"]
            title_level = entry["level"]

            title_start = segment_counter
            sub_segments = self._split_paragraphs_fine_grained(
                section_text.strip(), title_text, title_index
            )

            for seg in sub_segments:
                seg["segment_index"] = segment_counter
                seg["word_count"] = len(seg["content"])
                all_segments.append(seg)
                segment_counter += 1

            title_end = segment_counter - 1

            titles.append(
                {
                    "title_index": title_index,
                    "level": title_level,
                    "title": title_text,
                    "start_segment_index": title_start,
                    "end_segment_index": title_end,
                }
            )

        if not all_segments:
            all_segments.append(
                {
                    "segment_index": 0,
                    "title": "",
                    "title_index": -1,
                    "content": text.strip(),
                    "word_count": len(text.strip()),
                }
            )

        all_segments = self._merge_short_segments(all_segments)

        all_segments, titles = self._reindex_after_merge(all_segments, titles)

        return {"titles": titles, "segments": all_segments}

    def _merge_short_segments(
        self, segments: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """
        合并过短的 segment 到相邻 segment

        优先向后合并到同 title_index 的前一个 segment，
        如果是第一个 segment 则向前合并到下一个 segment。
        仅剩一个 segment 且过短时保留（不丢弃内容）。

        Args:
            segments: 原始 segment 列表

        Returns:
            合并后的 segment 列表
        """
        if len(segments) <= 1:
            return segments

        merged = [seg.copy() for seg in segments]
        i = 0

        while i < len(merged):
            if len(merged[i]["content"]) < self.MIN_SEGMENT_SIZE and len(merged) > 1:
                if i > 0:
                    target = i - 1
                else:
                    target = i + 1
                    i += 1

                merged[target]["content"] += "\n" + merged[i]["content"]
                merged[target]["word_count"] = len(merged[target]["content"])
                merged.pop(i)
            else:
                i += 1

        return merged

    def _reindex_after_merge(
        self, segments: List[Dict[str, Any]], titles: List[Dict[str, Any]]
    ) -> tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        """
        合并后重新编号 segment_index 并更新 title 的 start/end 范围

        Args:
            segments: 合并后的 segment 列表
            titles: 标题列表

        Returns:
            (重新编号的 segments, 更新范围的 titles)
        """
        for i, seg in enumerate(segments):
            seg["segment_index"] = i
            seg["word_count"] = len(seg["content"])

        seg_title_map: Dict[int, list[int]] = {}
        for i, seg in enumerate(segments):
            ti = seg["title_index"]
            seg_title_map.setdefault(ti, []).append(i)

        new_titles = []
        for title in titles:
            ti = title["title_index"]
            indices = seg_title_map.get(ti, [])
            if not indices:
                continue
            title["start_segment_index"] = indices[0]
            title["end_segment_index"] = indices[-1]
            new_titles.append(title)

        return segments, new_titles

    def _extract_title_structure(self, text: str) -> List[Dict[str, Any]]:
        """
        提取文本的标题层级结构

        使用 markdown-it 解析文本，正确识别 Markdown 标题并映射内容

        Args:
            text: 原始文本

        Returns:
            标题区域列表 [{title_index, level, title, content}, ...]
        """
        normalized = self._normalize_headings(text)
        md = markdown_it.MarkdownIt()
        tokens = md.parse(normalized)

        headings = []
        for i, token in enumerate(tokens):
            if token.type == "heading_open" and token.map:
                level = int(token.tag[1])
                inline_token = tokens[i + 1]
                heading_text = (
                    inline_token.content if inline_token.type == "inline" else ""
                )
                headings.append(
                    {
                        "level": level,
                        "title": heading_text,
                        "start_line": token.map[0],
                        "end_line": token.map[1],
                    }
                )

        lines = normalized.split("\n")
        sections = []
        title_counter = 0

        if headings and headings[0]["start_line"] > 0:
            pre_content = "\n".join(lines[: headings[0]["start_line"]]).strip()
            if pre_content:
                sections.append(
                    {
                        "title_index": -1,
                        "level": 1,
                        "title": "",
                        "content": pre_content,
                    }
                )

        for idx, heading in enumerate(headings):
            content_start = heading["end_line"]
            if idx + 1 < len(headings):
                content_end = headings[idx + 1]["start_line"]
            else:
                content_end = len(lines)

            content = "\n".join(lines[content_start:content_end]).strip()

            sections.append(
                {
                    "title_index": title_counter,
                    "level": heading["level"],
                    "title": heading["title"],
                    "content": content,
                }
            )
            title_counter += 1

        if not sections:
            sections.append(
                {
                    "title_index": 0,
                    "level": 1,
                    "title": "",
                    "content": text.strip(),
                }
            )

        return sections

    def _normalize_headings(self, text: str) -> str:
        """将非标准标题格式（中文序号、全大写行等）转换为 Markdown # 标题"""
        lines = text.split("\n")
        result = []
        in_code_block = False

        for line in lines:
            stripped = line.strip()

            if stripped.startswith("```") or stripped.startswith("~~~"):
                in_code_block = not in_code_block
                result.append(line)
                continue

            if in_code_block:
                result.append(line)
                continue

            if len(stripped) > 100 or stripped.startswith("|"):
                result.append(line)
                continue

            matched = False
            for pattern, pattern_type in self.TITLE_PATTERNS:
                match = pattern.match(stripped)
                if not match:
                    continue

                if pattern_type == "md":
                    result.append(line)
                    matched = True
                    break

                if pattern_type == "cn_num":
                    cn_char = match.group(1)
                    level = self._cn_char_to_level(cn_char, default=2)
                    result.append(f"{'#' * level} {stripped}")
                    matched = True
                    break

                if pattern_type == "num":
                    result.append(f"## {stripped}")
                    matched = True
                    break

                if pattern_type == "cn_paren":
                    cn_char = match.group(1)
                    level = self._cn_char_to_level(cn_char, default=2) + 1
                    result.append(f"{'#' * level} {stripped}")
                    matched = True
                    break

                if pattern_type == "num_paren":
                    result.append(f"### {stripped}")
                    matched = True
                    break

                if pattern_type == "chapter":
                    result.append(f"# {stripped}")
                    matched = True
                    break

            if not matched:
                if stripped.isupper() and len(stripped) < 50:
                    result.append(f"## {stripped}")
                else:
                    result.append(line)

        return "\n".join(result)

    def _cn_char_to_level(self, cn_char: str, default: int = 2) -> int:
        """
        中文数字字符映射到标题级别

        Args:
            cn_char: 中文数字字符
            default: 无法识别时的默认级别

        Returns:
            标题级别
        """
        cn_num_map = {
            "一": 1,
            "二": 2,
            "三": 3,
            "四": 4,
            "五": 5,
            "六": 6,
            "七": 7,
            "八": 8,
            "九": 9,
            "十": 10,
        }
        num = cn_num_map.get(cn_char)
        if num is None:
            return default
        if num <= 3:
            return 1
        elif num <= 6:
            return 2
        else:
            return 3

    def _split_paragraphs_fine_grained(
        self, content: str, title_text: str, title_index: int
    ) -> List[Dict[str, Any]]:
        """
        在标题区域内进行细粒度段落切分

        按双换行符拆分段落，超过 MAX_SEGMENT_SIZE 的段落按单换行进一步切分

        Args:
            content: 标题区域内容
            title_text: 所属标题文本
            title_index: 所属标题序号

        Returns:
            段落列表 [{title, title_index, content}, ...]
        """
        if not content:
            return []

        paragraphs = re.split(r"\n\s*\n", content)
        segments = []

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            if len(para) <= self.MAX_SEGMENT_SIZE:
                segments.append(
                    {
                        "title": title_text,
                        "title_index": title_index,
                        "content": para,
                    }
                )
            else:
                sub_segments = self._split_by_single_newline(
                    para, title_text, title_index
                )
                segments.extend(sub_segments)

        return (
            segments
            if segments
            else [
                {
                    "title": title_text,
                    "title_index": title_index,
                    "content": content,
                }
            ]
        )

    def _split_by_single_newline(
        self, content: str, title_text: str, title_index: int
    ) -> List[Dict[str, Any]]:
        """
        按单换行切分超长段落

        Args:
            content: 超长段落内容
            title_text: 所属标题文本
            title_index: 所属标题序号

        Returns:
            切分后的段落列表
        """
        lines = content.split("\n")
        segments = []
        current_chunk = []
        current_length = 0

        for line in lines:
            line_len = len(line)

            if current_length + line_len > self.MAX_SEGMENT_SIZE and current_chunk:
                segments.append(
                    {
                        "title": title_text,
                        "title_index": title_index,
                        "content": "\n".join(current_chunk),
                    }
                )
                current_chunk = [line]
                current_length = line_len
            else:
                current_chunk.append(line)
                current_length += line_len + 1

        if current_chunk:
            segments.append(
                {
                    "title": title_text,
                    "title_index": title_index,
                    "content": "\n".join(current_chunk),
                }
            )

        return (
            segments
            if segments
            else [
                {
                    "title": title_text,
                    "title_index": title_index,
                    "content": content,
                }
            ]
        )

    # ---- xlsx 按行分段 ----

    XLSX_MAX_SEGMENT_SIZE = 1000

    def segment_xlsx_by_row(self, text: str) -> Dict[str, Any]:
        """
        xlsx 专用分段：多行数据合并为一个 segment，超过 XLSX_MAX_SEGMENT_SIZE 时拆分

        _extract_from_xlsx() 输出中 Sheet 名称与表格数据以 \\n\\n 分隔，
        因此按 \\n\\n 拆分后需要相邻 block 配对（偶数=标题，奇数=表格数据）。

        每个 segment 的 content = 表头 + 分隔符 + 连续数据行。
        空行跳过，空 sheet 跳过，单行超长不拆分，不加上下文覆盖。
        返回结构与 smart_segment() 一致。

        Args:
            text: _extract_from_xlsx() 输出的 markdown 文本

        Returns:
            {"titles": [...], "segments": [...]}
        """
        sheet_blocks = re.split(r"\n\s*\n", text.strip())
        titles = []
        all_segments = []
        title_counter = 0
        segment_counter = 0

        i = 0
        while i < len(sheet_blocks):
            block = sheet_blocks[i].strip()
            if not block:
                i += 1
                continue

            title_match = re.match(r"^##\s+(.+)$", block)
            if title_match:
                sheet_name = title_match.group(1).strip()
                i += 1
                if i < len(sheet_blocks):
                    table_block = sheet_blocks[i].strip()
                    lines = table_block.split("\n")
                    if len(lines) >= 3:
                        header_line = lines[0].strip()
                        separator_line = lines[1].strip()
                        prefix = f"{header_line}\n{separator_line}"
                        prefix_len = len(prefix) + 1

                        title_index = title_counter
                        title_start = segment_counter

                        current_rows = []
                        current_length = prefix_len

                        for data_line in lines[2:]:
                            data_line = data_line.strip()
                            if not data_line:
                                continue

                            cells = [c.strip() for c in data_line.split("|")]
                            cells = [c for c in cells if c]
                            if not cells:
                                continue

                            row_len = len(data_line) + 1

                            if (
                                current_rows
                                and current_length + row_len
                                > self.XLSX_MAX_SEGMENT_SIZE
                            ):
                                content = prefix + "\n" + "\n".join(current_rows)
                                all_segments.append(
                                    {
                                        "segment_index": segment_counter,
                                        "title": sheet_name,
                                        "title_index": title_index,
                                        "content": content,
                                        "word_count": len(content),
                                    }
                                )
                                segment_counter += 1
                                current_rows = []
                                current_length = prefix_len

                            current_rows.append(data_line)
                            current_length += row_len

                        if current_rows:
                            content = prefix + "\n" + "\n".join(current_rows)
                            all_segments.append(
                                {
                                    "segment_index": segment_counter,
                                    "title": sheet_name,
                                    "title_index": title_index,
                                    "content": content,
                                    "word_count": len(content),
                                }
                            )
                            segment_counter += 1

                        title_end = segment_counter - 1

                        if segment_counter > title_start:
                            titles.append(
                                {
                                    "title_index": title_index,
                                    "level": 2,
                                    "title": sheet_name,
                                    "start_segment_index": title_start,
                                    "end_segment_index": title_end,
                                }
                            )
                            title_counter += 1
            i += 1

        if not all_segments:
            return {"titles": [], "segments": []}

        return {"titles": titles, "segments": all_segments}

    # ---- 兼容旧模式 ----

    def smart_segment_legacy(self, text: str) -> List[Dict[str, Any]]:
        """
        旧版智能分段（兼容已上传文档）

        按标题和段落分割文本，支持上下文重叠

        Args:
            text: 原始文本

        Returns:
            分段列表，每段包含 index, title, content, word_count
        """
        OVERLAP_SIZE = 100
        MAX_SEGMENT_SIZE = 2000

        raw_segments = self._split_by_structure_legacy(text)

        final_segments = []
        for i, seg in enumerate(raw_segments):
            if len(seg["content"]) > MAX_SEGMENT_SIZE:
                sub_segments = self._split_long_segment_legacy(
                    seg["content"], seg["title"]
                )
                final_segments.extend(sub_segments)
            else:
                final_segments.append(seg)

        segments_with_overlap = self._add_overlap(final_segments, OVERLAP_SIZE)

        for i, seg in enumerate(segments_with_overlap):
            seg["segment_index"] = i
            seg["word_count"] = len(seg["content"])

        return segments_with_overlap

    def _split_by_structure_legacy(self, text: str) -> List[Dict[str, Any]]:
        """按结构（标题、段落）分割文本（旧版）"""
        lines = text.split("\n")
        segments = []
        current_title = ""
        current_content = []

        for line in lines:
            stripped = line.strip()

            if not stripped:
                if current_content:
                    current_content.append("")
                continue

            is_title = self._is_title_line_legacy(stripped)

            if is_title:
                if current_content:
                    content = "\n".join(current_content).strip()
                    if content:
                        segments.append({"title": current_title, "content": content})
                current_title = stripped.lstrip("#").strip()
                current_content = [stripped]
            else:
                current_content.append(line)

        if current_content:
            content = "\n".join(current_content).strip()
            if content:
                segments.append({"title": current_title, "content": content})

        if not segments:
            segments.append({"title": "", "content": text.strip()})

        return segments

    def _is_title_line_legacy(self, line: str) -> bool:
        """判断是否为标题行（旧版）"""
        if len(line) > 100:
            return False

        flat_patterns = [
            re.compile(r"^#{1,6}\s+.+"),
            re.compile(r"^[一二三四五六七八九十]+[、.．].+"),
            re.compile(r"^\d+[、.．]\s*.+"),
            re.compile(r"^\([一二三四五六七八九十]+\).+"),
            re.compile(r"^\(\d+\).+"),
            re.compile(r"^第[一二三四五六七八九十\d]+[章节条].+"),
        ]
        for pattern in flat_patterns:
            if pattern.match(line):
                return True

        if line.startswith("|"):
            return False

        if line.isupper() and len(line) < 50:
            return True

        return False

    def _split_long_segment_legacy(
        self, content: str, title: str
    ) -> List[Dict[str, Any]]:
        """将过长的分段拆分（旧版）"""
        MAX_SEGMENT_SIZE = 2000
        segments = []
        paragraphs = re.split(r"\n\s*\n", content)

        current_chunk = []
        current_length = 0

        for para in paragraphs:
            para_length = len(para)

            if current_length + para_length > MAX_SEGMENT_SIZE and current_chunk:
                segments.append({"title": title, "content": "\n\n".join(current_chunk)})
                current_chunk = [para]
                current_length = para_length
            else:
                current_chunk.append(para)
                current_length += para_length + 2

        if current_chunk:
            segments.append({"title": title, "content": "\n\n".join(current_chunk)})

        return segments if segments else [{"title": title, "content": content}]

    def _add_overlap(
        self, segments: List[Dict[str, Any]], overlap_size: int
    ) -> List[Dict[str, Any]]:
        """为分段添加上下文重叠（旧版）"""
        if len(segments) <= 1:
            return segments

        result = []

        for i, seg in enumerate(segments):
            content = seg["content"]
            title = seg["title"]

            if i > 0:
                prev_content = segments[i - 1]["content"]
                overlap = self._get_overlap_text(
                    prev_content, overlap_size, from_end=True
                )
                if overlap:
                    content = f"[上文摘要]\n{overlap}\n\n{content}"

            if i < len(segments) - 1:
                next_content = segments[i + 1]["content"]
                overlap = self._get_overlap_text(
                    next_content, overlap_size, from_end=False
                )
                if overlap:
                    content = f"{content}\n\n[下文预览]\n{overlap}"

            result.append({"title": title, "content": content})

        return result

    def _get_overlap_text(self, text: str, size: int, from_end: bool = True) -> str:
        """获取重叠文本"""
        if len(text) <= size:
            return text

        if from_end:
            overlap = text[-size:]
            newline_pos = overlap.find("\n")
            if newline_pos > 0:
                overlap = overlap[newline_pos + 1 :]
            return overlap.strip()
        else:
            overlap = text[:size]
            newline_pos = overlap.rfind("\n")
            if newline_pos > 0:
                overlap = overlap[:newline_pos]
            return overlap.strip()

    # ---- 文件操作 ----

    async def save_file(self, file: UploadFile, knowledge_base_id: int) -> str:
        """
        保存上传的文件到本地

        Args:
            file: 上传的文件
            knowledge_base_id: 知识库ID

        Returns:
            文件存储路径
        """
        content = await file.read()

        filename = file.filename or "unknown"
        ext = self._get_file_extension(filename)

        content_hash = hashlib.md5(content).hexdigest()[:8]
        safe_filename = f"{knowledge_base_id}_{content_hash}.{ext}"

        dir_path = self._upload_dir / str(knowledge_base_id)
        dir_path.mkdir(parents=True, exist_ok=True)

        file_path = dir_path / safe_filename

        await asyncio.to_thread(_write_bytes, file_path, content)

        return str(file_path)


document_processor = DocumentProcessor()
